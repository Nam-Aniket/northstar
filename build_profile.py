#!/usr/bin/env python3
"""build_profile.py — Generate skills.json from a résumé file.

Usage:
    python build_profile.py --resume path/to/resume.docx
    python build_profile.py --resume path/to/resume.md --out my_skills.json
    python build_profile.py --text "Built dashboards in Power BI and SQL"
    echo "..." | python build_profile.py
    python build_profile.py --resume resume.docx --llm   # optional LLM pass
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from pathlib import Path

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    try:
        import pypdf  # noqa: F401
    except ImportError:
        print(
            "ERROR: pypdf not installed — cannot read PDF.\n"
            "  Install it:  pip install pypdf\n"
            "  Or paste résumé text via --text or stdin, or use a .docx file.",
            file=sys.stderr,
        )
        sys.exit(1)
    import pypdf

    reader = pypdf.PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(path: Path) -> str:
    """Return plain text from .docx, .pdf, .md, or .txt."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".md", ".txt", ""):
        return path.read_text(encoding="utf-8")
    # Try reading as plain text for unknown extensions
    return path.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# Deterministic matching
# ---------------------------------------------------------------------------

def _load_taxonomy(taxonomy_path: Path) -> dict:
    with taxonomy_path.open(encoding="utf-8") as f:
        return json.load(f)


def match_resume(text: str, taxonomy: dict) -> set[str]:
    """Return the set of skill labels whose aliases appear in text.

    Uses ontology.match_text (Aho-Corasick, full ontology) when available,
    with the curated taxonomy as the bundled fallback.  Falls back to the
    original per-label regex scan if ontology cannot be imported.
    """
    try:
        from ontology import match_text as _onto_match
        return _onto_match(text)
    except Exception:
        pass

    # Original fallback: per-label regex scan against the curated taxonomy
    from config import _alias_pattern
    lowered = text.lower()
    present: set[str] = set()
    for label, meta in taxonomy.items():
        for alias in meta["aliases"]:
            if _alias_pattern(alias).search(lowered):
                present.add(label)
                break
    return present


# ---------------------------------------------------------------------------
# Optional LLM pass
# ---------------------------------------------------------------------------

def _llm_match(text: str, taxonomy: dict) -> tuple[set[str], list[str]]:
    """
    Ask an LLM for skills, map them to taxonomy labels.
    Returns (extra_labels, unmatched_items).
    On any failure, returns (set(), []) and prints a warning.
    """
    try:
        from dotenv_loader import load_env
        load_env()
    except Exception:
        pass

    import os
    api_key = os.environ.get("LLM_API_KEY")
    if not api_key:
        print("WARNING: --llm requested but LLM_API_KEY not set — skipping LLM pass.", file=sys.stderr)
        return set(), []

    try:
        import urllib.request as _req
        import urllib.error as _err

        prompt = (
            "Return ONLY a JSON array of the concrete skills, tools, and technologies "
            "this person demonstrably has, based on their résumé text below.\n\n"
            + text[:6000]
        )
        payload = json.dumps({
            "model": "claude-3-haiku-20240307",
            "max_tokens": 512,
            "messages": [{"role": "user", "content": prompt}],
        }).encode()
        req = _req.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            method="POST",
        )
        with _req.urlopen(req, timeout=30) as resp:
            body = json.loads(resp.read())
        raw_content = body["content"][0]["text"].strip()
        # Extract JSON array from the response
        start = raw_content.find("[")
        end = raw_content.rfind("]") + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON array in LLM response")
        llm_items: list[str] = json.loads(raw_content[start:end])
    except Exception as exc:
        print(f"WARNING: LLM call failed ({exc}) — continuing with deterministic results.", file=sys.stderr)
        return set(), []

    # Map LLM items back to taxonomy labels
    # Build a reverse lookup: lowercased alias -> label
    alias_to_label: dict[str, str] = {}
    for label, meta in taxonomy.items():
        for alias in meta["aliases"]:
            alias_to_label[alias.lower()] = label
        alias_to_label[label.lower()] = label

    extra: set[str] = set()
    unmatched: list[str] = []
    for item in llm_items:
        key = item.lower().strip()
        if key in alias_to_label:
            extra.add(alias_to_label[key])
        else:
            unmatched.append(item)
    return extra, unmatched


# ---------------------------------------------------------------------------
# Output building
# ---------------------------------------------------------------------------

def build_skills_json(present: set[str], taxonomy: dict) -> dict:
    """Build the skills.json dict from present labels and full taxonomy.

    Taxonomy labels are partitioned into supported/unsupported as before.
    Labels in *present* that are NOT in the curated taxonomy (i.e. came from
    an extended ontology like ESCO) are appended to supported_skills using
    ontology.group_for() for their group.  unsupported_skills contains only
    curated taxonomy labels not matched — the full ontology label universe is
    never materialised there.
    """
    supported: dict = {}
    unsupported: dict = {}

    # Partition curated taxonomy labels
    for label in sorted(taxonomy.keys()):
        meta = taxonomy[label]
        if label in present:
            entry: dict = {
                "aliases": meta["aliases"],
                "group": meta["group"],
            }
            if meta.get("soft"):
                entry["soft"] = True
            supported[label] = entry
        else:
            unsupported[label] = meta["aliases"]

    # Emit ontology-only labels (not in curated taxonomy)
    try:
        from ontology import group_for as _group_for
    except Exception:
        def _group_for(lbl: str) -> str:  # type: ignore[misc]
            return "General"

    for label in sorted(present - set(taxonomy.keys())):
        supported[label] = {
            "aliases": [label.lower()],
            "group": _group_for(label),
        }

    return {"supported_skills": supported, "unsupported_skills": unsupported}


# ---------------------------------------------------------------------------
# Summary printer
# ---------------------------------------------------------------------------

def _print_summary(present: set[str], taxonomy: dict, out_path: Path) -> None:
    by_group: dict[str, list[str]] = defaultdict(list)
    for label in sorted(present):
        group = taxonomy[label]["group"]
        by_group[group].append(label)

    total = len(taxonomy)
    found = len(present)
    gaps = total - found
    groups = len(by_group)

    print(f"\nFound {found} skills on your résumé across {groups} groups -> {out_path}")
    print("REVIEW supported_skills: add any we missed, remove false positives.")
    print(f"({gaps} taxonomy skills not on your résumé are tracked as gaps.)\n")

    for group in sorted(by_group.keys()):
        labels = by_group[group]
        print(f"  [{group}]")
        for lbl in labels:
            print(f"    - {lbl}")
    print()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate skills.json from your résumé. Reads .docx, .md, .txt; .pdf needs pypdf."
    )
    source = parser.add_mutually_exclusive_group()
    source.add_argument("--resume", metavar="PATH", help="Path to résumé file (.docx/.md/.txt/.pdf)")
    source.add_argument("--text", metavar="TEXT", help='Résumé as a plain-text string (quote it)')
    parser.add_argument(
        "--taxonomy",
        metavar="PATH",
        default=str(Path(__file__).resolve().parent / "taxonomy.json"),
        help="Taxonomy file (default: taxonomy.json next to this script)",
    )
    parser.add_argument(
        "--out",
        metavar="PATH",
        default="skills.json",
        help="Output path (default: skills.json)",
    )
    parser.add_argument(
        "--llm",
        action="store_true",
        default=False,
        help="Run an optional LLM pass to catch skills the deterministic matcher missed (needs LLM_API_KEY)",
    )
    args = parser.parse_args()

    # --- Get résumé text ---
    if args.text:
        resume_text = args.text
    elif args.resume:
        resume_path = Path(args.resume)
        if not resume_path.exists():
            print(f"ERROR: File not found: {resume_path}", file=sys.stderr)
            sys.exit(1)
        resume_text = extract_text(resume_path)
    elif not sys.stdin.isatty():
        resume_text = sys.stdin.read()
    else:
        parser.print_help()
        sys.exit(1)

    if not resume_text.strip():
        print("ERROR: No résumé text found.", file=sys.stderr)
        sys.exit(1)

    # --- Load taxonomy ---
    taxonomy_path = Path(args.taxonomy)
    if not taxonomy_path.exists():
        print(f"ERROR: Taxonomy file not found: {taxonomy_path}", file=sys.stderr)
        sys.exit(1)
    taxonomy = _load_taxonomy(taxonomy_path)

    # --- Deterministic match ---
    present = match_resume(resume_text, taxonomy)

    # --- Optional LLM pass ---
    if args.llm:
        extra, unmatched = _llm_match(resume_text, taxonomy)
        present |= extra
        if unmatched:
            print("Consider adding these to taxonomy.json (LLM found them but they had no match):")
            for item in unmatched:
                print(f"  - {item}")

    # --- Build and write output ---
    out_path = Path(args.out)
    skills = build_skills_json(present, taxonomy)
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(skills, f, indent=2)
        f.write("\n")
    print(f"Wrote {out_path}")

    # --- Summary ---
    _print_summary(present, taxonomy, out_path)


if __name__ == "__main__":
    main()

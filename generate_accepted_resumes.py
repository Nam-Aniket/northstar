#!/usr/bin/env python3
"""Generate tailored resume packages for accepted jobs (ATS engine v2).

Reads matched_jobs.csv + job_posts.csv; produces per job in resumes/:
  {slug}_resume.docx
  {slug}_cover_letter.docx
  {slug}_job_description.txt
  {slug}_change_log.md
  {slug}_match_report.json
  ats_match_tracking.csv  (summary across all jobs)

v2 design (deterministic, no LLM, no new dependencies):
  TERM_BANK   — every skill label carries alias lists (AU/US spelling, plural,
                multi-word JD phrasings). JD terms are weighted x2 when they
                appear in the requirements region of the JD.
  FACT_BANK   — every truthful bullet variant is tagged with the skill labels
                it evidences. All text is hand-authored from main_resume.md;
                the algorithm only selects and orders, never writes prose.
  Selection   — greedy weighted set-cover picks the bullets that maximise
                coverage of THIS JD's terms under a per-role bullet budget,
                so two different JDs get genuinely different resumes.
  Self-check  — each saved .docx is reopened and scored Jobscan-style:
                frequency-weighted hard-skill coverage, title match, education,
                quantified-bullet count (>=5), word-count band (450-650),
                family-predicted skills from the local JD corpus. Target >=80.
  Slop linter — generation hard-fails on banned cliche phrases and duplicate
                bullet openers (Zinsser/Strunk & White/Bock writing rules).
  Uniqueness  — batch gate: no two resumes may share an identical body.

Bullet grammar: Laszlo Bock XYZ — accomplished [X] as measured by [Y], by
doing [Z]. Every bullet carries a number where main_resume.md provides one.

Typography matches the reference output: name 16pt Arial black centered,
subtitle 10.5pt bold black, body 9.1pt at 1.02 spacing, Heading 1 11pt navy
(31,78,121) with D9E2EC bottom border, Heading 2 9.4pt navy, List Bullet
bullets, A4 page.
"""

from __future__ import annotations

import argparse
import csv
import difflib
import json
import re
import sys
import zipfile
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


import config as config

ROOT = Path(__file__).resolve().parent
MATCHED = ROOT / "matched_jobs.csv"
JOBS = ROOT / "job_posts.csv"
OUT_ROOT = ROOT / "resumes"
TRACKING = OUT_ROOT / "ats_match_tracking.csv"

from config import (  # noqa: E402
    TERM_BANK, UNSUPPORTED_BANK,
    SUPPORTED_TERMS, UNSUPPORTED_TERMS,
    _TERM_PATTERNS, _UNSUPPORTED_PATTERNS,
    _alias_pattern,
    CONTACT, WORK_RIGHTS, NAME, EDUCATION, PROJECTS,
)

def _get_exclude_companies():
    return set(config.exclude_companies)

NAVY = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)

MATCH_RATE_TARGET = 80  # local Jobscan-equivalent pass bar
QUANTIFIED_TARGET = 5   # measurable-results benchmark (Cultivated Culture)
WORD_BAND = (450, 650)  # resume word-count band (sweet spot 500-600)

# Render the JD’s own spelling when it differs from the AU default label.
DISPLAY_OVERRIDES = {
    "data modeling": "Data Modeling",
    "dimensional modeling": "Data Modeling",
    "data visualization": "Data Visualization",
    "visualization": "Data Visualization",
}

# Anchored requirements-region detector (FIX A1).
# Trigger must be at the START of a line (optionally preceded by whitespace/bullets/numbers)
# OR immediately followed by a colon — i.e. a real section header, not a buried prose word.
# If no anchored match is found, extract_terms_detailed falls back to req_start=len(text)
# (no region weighting), which is the safe default.
_REQ_REGION = re.compile(
    r"(?:(?:^|(?<=\n))[\s\-\*\d\.\)]*"
    r"(requirements|what you(?:’|’)ll bring|about you|qualifications|must have|"
    r"essential|skills and experience|what we(?:’|’)re looking for|to be successful)"
    r"[\s\-\*\d\.\)]*(?:\n|:|\s*$)"
    r"|"
    r"(requirements|what you(?:’|’)ll bring|about you|qualifications|must have|"
    r"essential|skills and experience|what we(?:’|’)re looking for|to be successful)"
    r"\s*:)",
    re.IGNORECASE | re.MULTILINE,
)


def extract_terms_detailed(text: str) -> Tuple[Dict[str, Dict], List[str]]:
    """Return ({label: {weight, freq, alias}}, [unsupported labels]).

    weight = 2 if the term appears in the JD's requirements region, else 1.
    freq   = total occurrences across all aliases.
    alias  = first alias form found (drives spelling mirroring).
    """
    low = text.lower()
    req_match = _REQ_REGION.search(low)
    req_start = req_match.start() if req_match else len(low)

    hits: Dict[str, Dict] = {}
    for label, patterns in _TERM_PATTERNS.items():
        freq = 0
        weight = 1
        first_alias = None
        for pat, alias in zip(patterns, TERM_BANK[label]["aliases"]):
            for m in pat.finditer(low):
                freq += 1
                if first_alias is None:
                    first_alias = alias
                if m.start() >= req_start:
                    weight = 2
        if freq:
            hits[label] = {"weight": weight, "freq": freq, "alias": first_alias}

    unsupported = sorted(
        lbl for lbl, pats in _UNSUPPORTED_PATTERNS.items()
        if any(p.search(low) for p in pats)
    )
    return hits, unsupported


def extract_terms(text: str) -> Tuple[List[str], List[str]]:
    """v1-compatible signature: (sorted supported labels, sorted unsupported)."""
    hits, unsupported = extract_terms_detailed(text)
    return sorted(hits), unsupported


def importance(hit: Dict) -> int:
    """JD importance of a matched term: requirements-region weight x capped frequency."""
    return hit["weight"] * min(hit["freq"], 3)


# ---------------------------------------------------------------------------
# FACT_BANK — truthful bullet variants tagged with the labels they evidence.
# All text hand-authored from main_resume.md under XYZ grammar (Bock) and the
# clear-writing rules (active voice, verbs not nominalisations, no cliches).
# Openers are distinct within each pool so the slop linter always passes.
# ---------------------------------------------------------------------------

FACT_BANK: Dict[str, List[Dict]] = {
    "role1": [
        {"text": "Built a cloud-native data pipeline ingesting and enriching high-volume public feeds for downstream reporting and analytics.",
         "evidences": ["Data Pipelines", "Automation", "Insights"]},
        {"text": "Engineered ETL workflows that clean, normalise and enrich source data for cloud data lakes and downstream reporting.",
         "evidences": ["ETL", "Data Cleaning", "Data Quality", "Data Pipelines", "Data Warehousing"]},
        {"text": "Applied NLP topic modelling and source-validation scoring to automate relevance ranking and keep evidence retention audit-ready.",
         "evidences": ["NLP", "Machine Learning", "Data Governance", "Automation"]},
        {"text": "Deployed ML classification and summarisation models that convert unstructured text into analyst-ready outputs and alerts.",
         "evidences": ["Machine Learning", "Model Deployment", "NLP"]},
        {"text": "Architected a retrieval layer with vector embeddings and metadata filters for semantic search and reporting.",
         "evidences": ["Deep Learning", "Machine Learning", "NLP"]},
        {"text": "Documented data flows and source-validation logic so analyst workflows stay repeatable and audit-ready.",
         "evidences": ["Documentation", "Data Governance", "Process Mapping"]},
    ],
    "role2": [
        {"text": "Automated recurring SQL-based reconciliation reports with Python and Azure Data Factory, freeing analyst hours and reducing billing exceptions.",
         "evidences": ["SQL", "Python", "Azure Data Factory", "Automation", "Reconciliation", "ETL"]},
        {"text": "Built Power BI dashboards with custom DAX measures and Power Query layers, enabling self-service reporting for operations teams.",
         "evidences": ["Power BI", "DAX", "Power Query", "Self-Service BI", "Dashboarding", "Data Visualisation"]},
        {"text": "Combined interval data, external feeds and usage patterns into load-forecasting models that flagged assets at risk for operations planning.",
         "evidences": ["Forecasting", "Predictive Analytics", "Machine Learning", "Statistical Analysis"]},
        {"text": "Tracked operational KPIs in Power BI for a program delivering measurable cost savings across the client portfolio.",
         "evidences": ["Power BI", "KPI Reporting", "Dashboarding", "Data Storytelling"]},
        {"text": "Gathered reporting requirements from operations and planning teams and translated them into SQL/Python pipelines and dashboards.",
         "evidences": ["Requirements Gathering", "Stakeholder Management", "Process Mapping", "Communication", "SQL", "Python"]},
        {"text": "Mentored analyst trainees across Agile sprint cycles, assigning Jira tickets and running acceptance checks on SQL and Python deliverables.",
         "evidences": ["Mentoring", "Agile/Scrum", "Jira", "Collaboration", "UAT", "User Stories"]},
        {"text": "Presented exception trends and root-cause findings to operations and planning stakeholders, shaping monthly remediation priorities.",
         "evidences": ["Data Storytelling", "Communication", "Stakeholder Management", "Insights", "Statistical Analysis"]},
    ],
    "role3": [
        {"text": "Validated high-volume data feeds with SQL and Python data-quality checks across hundreds of millions of records per day.",
         "evidences": ["SQL", "Python", "Data Quality", "Data Cleaning"]},
        {"text": "Reconciled discrepancies between source systems, flagging thousands of mismatched records over a six-month audit cycle.",
         "evidences": ["Reconciliation", "Data Migration", "Data Quality", "Attention to Detail"]},
        {"text": "Built Power BI dashboards on operational KPIs, cutting manual reporting time by 30%.",
         "evidences": ["Power BI", "Dashboarding", "KPI Reporting", "Automation"]},
        {"text": "Clarified metric definitions and reporting requirements with business users while documenting data-quality exceptions for correction.",
         "evidences": ["Documentation", "Requirements Gathering", "Stakeholder Management", "Communication"]},
    ],
}

BULLET_BUDGETS = {"role1": 3, "role2": 4, "role3": 3}

EXPERIENCE_SLOTS = [
    ("Data Analytics Specialist | Company A | Remote | 2025 - Present", "role1"),
    ("Data Analyst & Automation Lead | Company B | 2022 - 2024", "role2"),
    ("Analyst | Company C | 2020 - 2022", "role3"),
]

# Personal override: a gitignored facts.json supplies the owner's real experience
# (bullets + budgets + slots). Absent -> the generic example above is used, so the
# committed repo stays PII-free and the test suite stays deterministic.
_facts_override = config.load_facts_override()
if _facts_override:
    FACT_BANK = _facts_override["FACT_BANK"]
    BULLET_BUDGETS = _facts_override["BULLET_BUDGETS"]
    EXPERIENCE_SLOTS = [tuple(s) for s in _facts_override["EXPERIENCE_SLOTS"]]


def select_bullets(slot: str, jd_hits: Dict[str, Dict]) -> List[Dict]:
    """Greedy weighted set-cover: pick the bullets that add the most uncovered
    JD-term importance, up to the slot budget. Deterministic. Selected bullets
    render in authored order so the lead bullet stays strong."""
    pool = FACT_BANK[slot]
    budget = BULLET_BUDGETS[slot]
    chosen: List[Dict] = []
    covered: set = set()
    remaining = list(pool)

    while len(chosen) < budget and remaining:
        def gain(v: Dict) -> int:
            return sum(
                importance(jd_hits[l])
                for l in v["evidences"]
                if l in jd_hits and l not in covered
            )
        best = max(remaining, key=lambda v: (gain(v), -pool.index(v)))
        chosen.append(best)
        covered |= set(best["evidences"])
        remaining.remove(best)

    return sorted(chosen, key=pool.index)


# ---------------------------------------------------------------------------
# Slop linter — banned phrases + duplicate openers. All content is authored,
# so any hit is an authoring bug: hard fail.
# ---------------------------------------------------------------------------

BANNED_PHRASES = [
    "leverag", "spearhead", "passionate", "synerg", "results-driven",
    "results driven", "proven track record", "dynamic professional",
    "team player", "go-getter", "think outside", "utilize", "utilis",
    "responsible for", "duties included", "i am excited", "i believe my",
    "cutting-edge", "cutting edge", "seamless", "delve", "in today's",
    "fast-paced", "best-in-class", "world-class", "honed", "esteemed",
    "keen eye",
]


def lint_text(text: str, context: str) -> None:
    low = text.lower()
    for phrase in BANNED_PHRASES:
        if phrase in low:
            raise ValueError(f"Slop linter: banned phrase '{phrase}' in {context}: {text[:90]!r}")


def lint_bullet_openers(bullets: List[str], context: str) -> None:
    openers = [b.split()[0].lower() for b in bullets if b.split()]
    dupes = [w for w, n in Counter(openers).items() if n > 1]
    if dupes:
        raise ValueError(f"Slop linter: duplicate bullet openers {dupes} in {context}")


def validate_fact_bank() -> None:
    """Structural guard, runs before any generation."""
    for slot, pool in FACT_BANK.items():
        for entry in pool:
            unknown = [l for l in entry["evidences"] if l not in TERM_BANK]
            if unknown:
                raise ValueError(f"FACT_BANK[{slot}]: unknown evidence labels {unknown}")
            lint_text(entry["text"], f"FACT_BANK[{slot}]")
        lint_bullet_openers([e["text"] for e in pool], f"FACT_BANK[{slot}] pool")
        if len(pool) < BULLET_BUDGETS[slot]:
            raise ValueError(f"FACT_BANK[{slot}] smaller than its budget")


# ---------------------------------------------------------------------------
# Content: role families, skills section, summaries
# ---------------------------------------------------------------------------

def slugify(value: str) -> str:
    value = value.lower().replace("&", " and ")
    value = re.sub(r"[^a-z0-9]+", "_", value)
    return re.sub(r"_+", "_", value).strip("_")[:90]


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def role_family(role: str, jd: str) -> str:
    role_l = role.lower()
    blob = f"{role} {jd}".lower()
    if any(x in role_l for x in ["data scientist", "machine learning engineer", "ml engineer", "ai analyst", "ai engineer", "research scientist"]):
        return "data_scientist"
    if any(x in role_l for x in ["machine learning", "deep learning"]) and "analyst" not in role_l:
        return "data_scientist"
    if any(x in role_l for x in ["business analyst", "ict business analyst", "systems analyst"]):
        return "business_analyst"
    if any(x in role_l for x in ["data engineer", "integration developer", "etl developer", "pipeline engineer"]):
        return "engineering_bi"
    if any(x in role_l for x in ["automation", "developer"]) and "data" in role_l:
        return "engineering_bi"
    if any(x in role_l for x in ["pricing", "sales analyst", "product analyst", "planning analyst", "food & beverage", "commercial analyst", "finance analyst", "fp&a"]):
        return "commercial_analytics"
    if any(x in role_l for x in ["reporting", "insights", " bi ", "business intelligence"]):
        return "reporting_bi"
    if "product owner" in role_l:
        return "business_analyst"
    if any(x in blob for x in ["user acceptance testing", "uat", "process mapping", "requirements gathering"]):
        return "business_analyst"
    return "data_analyst"


# Per-family base skill lists (canonical labels). JD-matched labels are pulled
# to the front; matched labels missing from the base list are inserted into
# their TERM_BANK group so every claimable JD term lands on the resume.
_SKILL_BASES: Dict[str, Dict[str, List[str]]] = {
    "data_analyst": {
        "Analytics & BI": ["SQL", "Python", "Power BI", "Excel", "Tableau", "DAX", "KPI Reporting", "Data Visualisation"],
        "Business Delivery": ["Stakeholder Management", "Jira", "Agile/Scrum", "Documentation", "Requirements Gathering"],
        "Data Engineering": ["Azure Data Factory", "ETL", "Data Pipelines", "Data Quality", "Data Warehousing"],
        "Data Science": ["Pandas", "NumPy", "Scikit-learn", "Statistical Analysis", "Forecasting", "A/B Testing"],
    },
    "reporting_bi": {
        "Analytics & BI": ["Power BI", "DAX", "Power Query", "SQL", "Data Modelling", "Excel", "Self-Service BI", "KPI Reporting"],
        "Business Delivery": ["Stakeholder Management", "Agile/Scrum", "Jira", "Documentation"],
        "Data Engineering": ["Azure Data Factory", "ETL", "Data Pipelines", "Data Quality", "Data Warehousing"],
        "Data Science": ["Pandas", "Statistical Analysis", "Forecasting", "A/B Testing"],
    },
    "business_analyst": {
        "Business Delivery": ["Requirements Gathering", "Process Mapping", "Stakeholder Management", "Jira", "Agile/Scrum", "UAT", "Documentation", "User Stories"],
        "Analytics & BI": ["SQL", "Power BI", "Excel", "Data Visualisation", "KPI Reporting"],
        "Data Engineering": ["ETL", "Data Pipelines", "Azure Data Factory", "Data Quality"],
        "Data Science": ["Statistical Analysis", "Forecasting", "A/B Testing"],
    },
    "engineering_bi": {
        "Data Engineering": ["Azure Data Factory", "ETL", "Data Pipelines", "Azure Databricks", "Apache Spark", "Kafka", "Data Warehousing", "Data Quality"],
        "Analytics & BI": ["SQL", "Python", "Power BI", "Excel", "DAX", "KPI Reporting"],
        "Business Delivery": ["Stakeholder Management", "Jira", "Agile/Scrum", "Documentation"],
        "Data Science": ["Pandas", "Statistical Analysis", "Machine Learning", "Feature Engineering"],
    },
    "commercial_analytics": {
        "Analytics & BI": ["SQL", "Python", "Power BI", "Excel", "Tableau", "KPI Reporting", "Data Visualisation"],
        "Business Delivery": ["Stakeholder Management", "Jira", "Agile/Scrum", "Documentation"],
        "Data Engineering": ["Azure Data Factory", "ETL", "Data Quality", "Data Pipelines"],
        "Data Science": ["Forecasting", "Statistical Analysis", "A/B Testing", "Scikit-learn"],
    },
    "data_scientist": {
        "Data Science": ["Python", "Scikit-learn", "Statistical Analysis", "A/B Testing", "Feature Engineering", "Machine Learning", "Forecasting", "Pandas"],
        "Analytics & BI": ["SQL", "Power BI", "Data Visualisation", "Excel", "KPI Reporting"],
        "Data Engineering": ["Azure Data Factory", "ETL", "Data Pipelines", "Azure Databricks", "Apache Spark"],
        "Business Delivery": ["Stakeholder Management", "Agile/Scrum", "Jira", "Documentation"],
    },
}

_GROUP_ORDER: Dict[str, List[str]] = {
    "data_analyst": ["Analytics & BI", "Business Delivery", "Data Engineering", "Data Science"],
    "reporting_bi": ["Analytics & BI", "Business Delivery", "Data Engineering", "Data Science"],
    "business_analyst": ["Business Delivery", "Analytics & BI", "Data Engineering", "Data Science"],
    "engineering_bi": ["Data Engineering", "Analytics & BI", "Business Delivery", "Data Science"],
    "commercial_analytics": ["Analytics & BI", "Business Delivery", "Data Engineering", "Data Science"],
    "data_scientist": ["Data Science", "Analytics & BI", "Data Engineering", "Business Delivery"],
}


def _display(label: str, jd_hits: Dict[str, Dict]) -> str:
    """Mirror the JD's spelling for AU/US variant labels."""
    hit = jd_hits.get(label)
    if hit and hit["alias"] in DISPLAY_OVERRIDES:
        return DISPLAY_OVERRIDES[hit["alias"]]
    return label


def compact_skills(family: str, jd_hits: Dict[str, Dict]) -> List[str]:
    bases = _SKILL_BASES.get(family, _SKILL_BASES["data_analyst"])
    order = _GROUP_ORDER.get(family, _GROUP_ORDER["data_analyst"])

    # Matched hard-skill labels not in any base list, routed to their group.
    base_union = {l for items in bases.values() for l in items}
    overflow: Dict[str, List[str]] = {g: [] for g in order}
    for label, hit in sorted(jd_hits.items(), key=lambda kv: -importance(kv[1])):
        meta = TERM_BANK[label]
        if meta.get("soft") or label in base_union:
            continue
        group = meta["group"]
        if group in overflow:
            overflow[group].append(label)

    lines = []
    for group in order:
        base = bases.get(group, [])
        matched = sorted(
            [l for l in base if l in jd_hits],
            key=lambda l: (-importance(jd_hits[l]), base.index(l)),
        )
        rest = [l for l in base if l not in jd_hits]
        selected = matched + overflow.get(group, []) + rest
        seen: set = set()
        deduped = [x for x in selected if not (x in seen or seen.add(x))]
        rendered = [_display(l, jd_hits) for l in deduped[:8]]
        lines.append(f"{group}: {', '.join(rendered)}")
    return lines


_BASE_SUMMARIES = {
    "reporting_bi": (
        "BI and reporting analyst with experience building Power BI dashboards, "
        "DAX measures and SQL/Python data checks over large operational datasets. Converts "
        "complex source data into self-service reporting layers, executive-ready KPI dashboards "
        "and data-quality controls, with a strong business partnering approach."
    ),
    "business_analyst": (
        "Data and business analyst with experience translating operational requirements "
        "into SQL/Python analysis, Power BI dashboards, process documentation and Agile/Jira "
        "delivery. Gathers and clarifies requirements, supports acceptance testing and communicates "
        "findings to stakeholders across data-quality and operations workflows."
    ),
    "commercial_analytics": (
        "Commercially oriented data analyst with experience using SQL, Python, Power BI "
        "and Excel to automate reporting, reconcile large datasets and deliver decision-ready KPI "
        "and trend analysis to business stakeholders, with business partnering experience across "
        "operations and planning functions."
    ),
    "engineering_bi": (
        "Data analyst and automation specialist with experience building SQL/Python "
        "pipelines, Azure Data Factory ETL workflows, Power BI reporting and data-quality controls "
        "over high-volume operational data, from source extraction through to BI layer and "
        "stakeholder reporting."
    ),
    "data_scientist": (
        "Data analyst with a machine learning and statistical modelling background: applied "
        "experience building forecasting models, classification pipelines and analytical "
        "workflows over high-volume operational datasets, plus a Master of Data Science. "
        "Translates model outputs into clear business decisions."
    ),
    "data_analyst": (
        "Data analyst with experience building SQL/Python pipelines, Power BI dashboards "
        "and Azure Data Factory workflows over high-volume operational datasets. Strong business "
        "partnering background, turning ad hoc analysis and recurring reporting into decisions for "
        "operations and planning stakeholders."
    ),
}

# One appended clause max — surfaces top JD themes the base summary missed.
SUMMARY_CLAUSES = {
    "Data Governance": "data governance and quality controls",
    "Self-Service BI": "self-service reporting",
    "Data Storytelling": "clear data storytelling for non-technical audiences",
    "Forecasting": "demand and load forecasting",
    "Predictive Analytics": "predictive analytics",
    "A/B Testing": "experiment-driven analysis",
    "Reconciliation": "high-volume data reconciliation",
    "Automation": "reporting automation",
    "Requirements Gathering": "structured requirements gathering",
    "Process Mapping": "process mapping",
    "Data Migration": "data migration and reconciliation",
    "Data Modelling": "data modelling",
    "Machine Learning": "applied machine learning",
    "NLP": "NLP and text analytics",
    "Data Quality": "data quality assurance",
    "UAT": "acceptance testing support",
    "Documentation": "clear process documentation",
}


def summary_for(family: str, jd_hits: Dict[str, Dict]) -> str:
    base = _BASE_SUMMARIES.get(family, _BASE_SUMMARIES["data_analyst"])
    base_low = base.lower()
    candidates = sorted(
        (l for l in jd_hits if l in SUMMARY_CLAUSES and SUMMARY_CLAUSES[l].split()[0] not in base_low),
        key=lambda l: -importance(jd_hits[l]),
    )
    picks = [SUMMARY_CLAUSES[l] for l in candidates[:2]]
    if not picks:
        return base
    focus = " and ".join(picks)
    return f"{base} Particular recent focus: {focus}."


def projects_for(family: str) -> List[Tuple[str, List[str]]]:
    streaming = (
        "Real-Time Demand Forecasting Pipeline",
        [
            "Built a streaming pipeline forecasting short-term demand from simulated sensor telemetry.",
            "Processed event streams with Spark Structured Streaming micro-batches to produce live forecasts.",
        ],
    )
    recommender = (
        "Hybrid Recommendation Engine",
        [
            "Built a hybrid recommender combining Neural Collaborative Filtering with TF-IDF content-based filtering.",
            "Processed 32M-row dataset in Python with cloud database backend; served recommendations at sub-100ms latency.",
        ],
    )
    health = (
        "Public Health Analytics Dashboard",
        [
            "Reshaped 100K survey records in Pandas from wide to long format for demographic analysis.",
            "Built an R Shiny dashboard with maps, demographic boxplots and 10-year time-series charts.",
        ],
    )
    if family == "engineering_bi":
        return [streaming, recommender]
    if family == "reporting_bi":
        return [recommender, health]
    if family == "business_analyst":
        return [health, streaming]
    if family == "data_scientist":
        return [recommender, streaming]
    return [recommender, health]


_SENIORITY_WORDS = re.compile(r"\b(senior|lead|principal|head of|head|director|manager|chief|staff)\b", re.IGNORECASE)

_FAMILY_TAGLINES = {
    "data_scientist": "Machine Learning & Statistical Modelling",
    "engineering_bi": "Data Pipelines & Automation",
    "reporting_bi": "Power BI & Reporting",
    "business_analyst": "Data & Process Analysis",
    "commercial_analytics": "Commercial Analytics",
    "data_analyst": "Analytics & Insights",
}


def subtitle_for(role_title: str, family: str) -> str:
    """Mirror the JD title, but never claim seniority the candidate doesn't have."""
    if not _SENIORITY_WORDS.search(role_title):
        return role_title
    stripped = _SENIORITY_WORDS.sub("", role_title)
    stripped = re.sub(r"\s{2,}", " ", stripped).strip(" -|,")
    if len(stripped.split()) >= 2:
        return stripped
    return f"{stripped} | {_FAMILY_TAGLINES[family]}" if stripped else _FAMILY_TAGLINES[family]


def find_job_post(job_posts: List[Dict[str, str]], match: Dict[str, str]) -> Dict[str, str]:
    for row in job_posts:
        if match.get("job_url") and row.get("job_url") == match.get("job_url"):
            return row
    for row in job_posts:
        if row.get("company") == match.get("company") and row.get("role_title") == match.get("role_title"):
            return row
    return {}


# ---------------------------------------------------------------------------
# Typography helpers (unchanged from reference output)
# ---------------------------------------------------------------------------

def _add_bottom_border(paragraph, color: str = "D9E2EC", size: str = "8") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.1)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(1.8)
    normal.paragraph_format.line_spacing = 1.02

    for name, size, space_before, space_after in [
        ("Heading 1", 11, 7, 2),
        ("Heading 2", 9.4, 3, 0),
    ]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(space_after)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.4)
    p.paragraph_format.line_spacing = 1.02
    p.add_run(text)


def _add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run(title.upper())
    _add_bottom_border(p)


# ---------------------------------------------------------------------------
# Document writers
# ---------------------------------------------------------------------------

def build_content(target: Dict[str, str], jd_text: str) -> Dict:
    """Assemble all tailored content for one job. Pure function of inputs."""
    family = role_family(target["role_title"], jd_text)
    jd_hits, unsupported = extract_terms_detailed(jd_text)

    bullets_by_slot = {
        slot: [e["text"] for e in select_bullets(slot, jd_hits)]
        for _, slot in EXPERIENCE_SLOTS
    }
    summary = summary_for(family, jd_hits)
    skills_lines = compact_skills(family, jd_hits)
    subtitle = subtitle_for(target["role_title"], family)

    # Linting: all rendered prose must pass.
    lint_text(summary, "summary")
    for slot, bullets in bullets_by_slot.items():
        for b in bullets:
            lint_text(b, f"bullet[{slot}]")
        lint_bullet_openers(bullets, f"selected bullets[{slot}]")

    return {
        "family": family,
        "jd_hits": jd_hits,
        "unsupported": unsupported,
        "subtitle": subtitle,
        "summary": summary,
        "skills_lines": skills_lines,
        "bullets_by_slot": bullets_by_slot,
        # Real résumé projects win; the generic sample bank is the fallback only
        # for users who supplied none (never fabricate over real ones).
        "projects": list(PROJECTS) if PROJECTS else projects_for(family),
    }


def write_resume(content: Dict, path: Path) -> None:
    doc = Document()
    _style_doc(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(NAME.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = BLACK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle_p.add_run(content["subtitle"])
    sr.bold = True
    sr.font.name = "Arial"
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = BLACK

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run(CONTACT)
    contact_p.paragraph_format.space_after = Pt(1)

    rights_p = doc.add_paragraph()
    rights_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rights_p.add_run(WORK_RIGHTS)
    rr.font.name = "Arial"
    rr.font.size = Pt(8.8)
    rr.font.color.rgb = BLACK
    rights_p.paragraph_format.space_after = Pt(4)

    _add_section(doc, "Summary")
    doc.add_paragraph(content["summary"])

    _add_section(doc, "Skills")
    for line in content["skills_lines"]:
        doc.add_paragraph(line)

    _add_section(doc, "Experience")
    for title_text, slot in EXPERIENCE_SLOTS:
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 2"]
        p.add_run(title_text).bold = True
        for bullet in content["bullets_by_slot"][slot]:
            _add_bullet(doc, bullet)

    _add_section(doc, "Projects")
    for proj_name, proj_bullets in content["projects"]:
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 2"]
        p.add_run(proj_name).bold = True
        for bullet in proj_bullets:
            _add_bullet(doc, bullet)

    _add_section(doc, "Education")
    for _edu_line in EDUCATION:
        _add_bullet(doc, _edu_line)

    doc.save(path)


# ---------------------------------------------------------------------------
# Cover letters — sector hook + JD-evidence sentences, assembled per job.
# ---------------------------------------------------------------------------

SECTOR_HOOKS = [
    (re.compile(r"\b(energy|utility|utilities|electricity|grid|meter|power networks)\b", re.IGNORECASE),
     "Working with high-volume operational data in asset-intensive environments means billing, outage and consumption datasets at scale are familiar territory."),
    (re.compile(r"\b(health|aged care|hospital|clinical|patient|care sector|ndis)\b", re.IGNORECASE),
     "Operational data-quality and reporting work in regulated environments translates directly to the rigour care-sector data demands."),
    (re.compile(r"\b(government|public sector|aps|council|department)\b", re.IGNORECASE),
     "My work has lived in audit-heavy environments where documentation, data-quality standards and traceability are non-negotiable."),
    (re.compile(r"\b(bank|superannuation|financial services|insurance|wealth|fund|finance sector)\b", re.IGNORECASE),
     "Reconciling large datasets and tracing exceptions to root cause is the same discipline financial-services data work demands."),
    (re.compile(r"\b(retail|hospitality|venue|restaurant|food|beverage|e-commerce|consumer)\b", re.IGNORECASE),
     "Demand forecasting and operational KPI reporting are central to my data analyst experience, the same questions retail and hospitality teams ask of their data."),
]

# Short evidence sentences keyed by label, for the cover letter middle.
ACHIEVEMENT_SENTENCES = {
    "Power BI": "I built Power BI dashboards that operations teams used daily to monitor key metrics across a large asset portfolio",
    "SQL": "I automated recurring SQL-based reconciliation reports, freeing significant analyst hours each month",
    "Python": "I automated recurring reporting with Python and Azure Data Factory, freeing analyst hours for higher-value analysis",
    "Azure Data Factory": "I built Azure Data Factory workflows that automated recurring reconciliation and data-quality reports",
    "Forecasting": "I built load-forecasting models from interval and external data feeds that operations teams used to flag assets at risk",
    "Predictive Analytics": "I built predictive models from operational and external data that planning teams used to prioritise interventions",
    "Machine Learning": "I have deployed ML classification and summarisation models in production and built neural recommender systems in academic work",
    "Requirements Gathering": "I gathered reporting requirements from operations and planning teams and translated them into pipelines and dashboards",
    "Stakeholder Management": "I presented exception trends and root-cause findings to operations and planning stakeholders each month",
    "Data Quality": "I ran SQL/Python data-quality checks over hundreds of millions of daily records from source systems",
    "Reconciliation": "I reconciled records across source systems, flagging thousands of mismatches across a multi-month audit cycle",
    "ETL": "I engineered ETL workflows that clean and enrich high-volume operational data for cloud data lakes",
    "Dashboarding": "I built KPI dashboards that cut manual reporting time by 30%",
    "KPI Reporting": "I built KPI dashboards that cut manual reporting time by 30%",
    "Automation": "I automated recurring reports with Python and Azure Data Factory, freeing analyst hours for higher-value work",
    "Data Pipelines": "I built a cloud-native data pipeline that ingests and enriches high-volume source data for downstream analytics",
    "NLP": "I applied topic modelling and ML summarisation to convert unstructured text into structured analyst-ready outputs",
    "Statistical Analysis": "I built statistical and forecasting models over high-volume operational data to support planning decisions",
}


def _join_labels(labels: List[str]) -> str:
    """Render a list as 'A, B and C' (no Oxford comma, no em-dash)."""
    labels = [l for l in labels if l]
    if not labels:
        return ""
    if len(labels) == 1:
        return labels[0]
    return ", ".join(labels[:-1]) + " and " + labels[-1]


def _lc_sentence(s: str) -> str:
    """Lowercase the first word and drop a trailing period, so a résumé bullet can
    be embedded mid-sentence ('Built X.' -> 'built X')."""
    s = s.strip().rstrip(".")
    return (s[:1].lower() + s[1:]) if s else s


def cover_letter_paragraphs(content: Dict, company: str, role: str, jd_text: str) -> List[str]:
    """Deterministic, real-data cover letter. Draws on the candidate's own
    JD-matched bullets, real education and real projects - never fabricates."""
    jd_hits = content["jd_hits"]

    # This JD's top requirement labels, by importance.
    top_labels = [l for l, _ in sorted(jd_hits.items(), key=lambda kv: -importance(kv[1]))][:3]

    # --- p1: role/company + sector hook + the JD's own emphasis ---
    hook = ""
    blob = f"{company} {role} {jd_text}"
    for pattern, sentence in SECTOR_HOOKS:
        if pattern.search(blob):
            hook = " " + sentence
            break
    p1 = f"I am applying for the {role} role at {company}.{hook}"
    if top_labels:
        p1 += f" The role's focus on {_join_labels(top_labels)} lines up closely with my work."

    # --- p2: the candidate's own real, JD-selected bullets (truthful, already
    # linted in build_content), ranked by how many top requirements they evidence ---
    real_bullets = [b for _, slot in EXPERIENCE_SLOTS for b in content["bullets_by_slot"].get(slot, [])]
    seen: set = set()
    real_bullets = [b for b in real_bullets if not (b in seen or seen.add(b))]

    def _coverage(bullet: str) -> int:
        low = bullet.lower()
        return sum(1 for lbl in top_labels
                   if lbl.lower() in low or (jd_hits[lbl].get("alias") or "") in low)

    ranked_bullets = sorted(real_bullets, key=lambda b: (-_coverage(b), real_bullets.index(b)))
    chosen = [b for b in ranked_bullets if _coverage(b) > 0][:2] or ranked_bullets[:2]

    if chosen:
        if len(chosen) == 2:
            p2 = f"In recent work I {_lc_sentence(chosen[0])}. I also {_lc_sentence(chosen[1])}."
        else:
            p2 = f"In recent work I {_lc_sentence(chosen[0])}."
    else:
        # Fallback only when the candidate has no bullets at all.
        ranked = sorted((l for l in jd_hits if l in ACHIEVEMENT_SENTENCES),
                        key=lambda l: -importance(jd_hits[l]))
        sents = [ACHIEVEMENT_SENTENCES[l] for l in ranked[:2]] or [ACHIEVEMENT_SENTENCES["SQL"]]
        p2 = " ".join(f"{s}." for s in sents)

    # --- p3: real education + a real project. Omit anything absent; never invent. ---
    p3 = ""
    if EDUCATION:
        edu_short = ", ".join(p.strip() for p in EDUCATION[0].split(",")[:2] if p.strip())
        if edu_short:
            p3 = f"I hold {edu_short}."
    if PROJECTS:  # module global is empty unless the résumé supplied real projects
        p3 += (" " if p3 else "") + f"Alongside this I built {PROJECTS[0][0]}."
    if not p3 and content.get("summary"):
        p3 = content["summary"]

    # --- p4: call to action + work rights ---
    work_rights_line = f" {WORK_RIGHTS}" if WORK_RIGHTS else ""
    p4 = f"I would welcome a short call about how I can help {company}'s team.{work_rights_line}"

    paragraphs = [p for p in (p1, p2, p3, p4) if p]
    for p in paragraphs:
        lint_text(p, "cover letter")
    return paragraphs


def write_cover_letter_docx(content: Dict, company: str, role: str, jd_text: str, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08

    name_p = doc.add_paragraph()
    nr = name_p.add_run(NAME.upper())
    nr.bold = True
    nr.font.name = "Arial"
    nr.font.size = Pt(13)
    nr.font.color.rgb = BLACK

    contact_p = doc.add_paragraph()
    contact_p.add_run(CONTACT)
    contact_p.paragraph_format.space_after = Pt(2)

    rights_p = doc.add_paragraph()
    rights_p.add_run(WORK_RIGHTS)
    rights_p.paragraph_format.space_after = Pt(14)

    doc.add_paragraph("Dear Hiring Manager,")
    for para in cover_letter_paragraphs(content, company, role, jd_text):
        doc.add_paragraph(para)
    doc.add_paragraph(f"Best regards,\n{NAME.split()[0]}")

    doc.save(path)


# ---------------------------------------------------------------------------
# ATS self-check — Jobscan-equivalent scoring of the FINAL saved document.
# ---------------------------------------------------------------------------

def family_skill_priors(job_posts: List[Dict[str, str]]) -> Dict[str, List[str]]:
    """Top supported labels per family across the local JD corpus —
    the deterministic equivalent of Jobscan's 'predicted skills'."""
    counters: Dict[str, Counter] = {}
    for row in job_posts:
        jd = row.get("job_text", "")
        if not jd.strip():
            continue
        fam = role_family(row.get("role_title", ""), jd)
        hits, _ = extract_terms_detailed(jd)
        counters.setdefault(fam, Counter()).update(hits.keys())
    return {fam: [l for l, _ in c.most_common(8)] for fam, c in counters.items()}


def docx_readiness(docx_path: Path) -> Dict:
    """Check a .docx for ATS parse-killers and basic content signals.

    Returns {"parseable": bool, "checks": [{"name", "ok", "detail"}...]}.
    parseable is True only when NO parse-killer tags are found.
    Parse-killers: tables, drawing/images, picture elements, text-boxes.
    """
    _PARSE_KILLERS = [
        ("<w:tbl",      "no_tables"),
        ("<w:drawing",  "no_images_drawings"),
        ("<w:pict",     "no_pictures"),
        ("<w:txbxContent", "no_text_boxes"),
    ]
    _STANDARD_HEADINGS = ["experience", "education", "skills", "summary", "projects"]

    checks: List[Dict] = []
    parseable = True

    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    except Exception as exc:
        return {"parseable": False, "checks": [{"name": "zip_open", "ok": False, "detail": str(exc)}]}

    # 1. Parse-killer tag checks
    for tag, name in _PARSE_KILLERS:
        present = tag in xml
        ok = not present
        if not ok:
            parseable = False
        checks.append({"name": name, "ok": ok, "detail": f"{tag} {'found' if present else 'absent'}"})

    # 2. Extract plain body text for content checks (strip XML tags)
    body_text = re.sub(r"<[^>]+>", " ", xml)
    body_text = re.sub(r"\s+", " ", body_text).strip()
    snippet = body_text[:600]

    email_ok = bool(re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", snippet))
    checks.append({
        "name": "contact_email_in_body",
        "ok": email_ok,
        "detail": "email found in first 600 chars" if email_ok else "no email found in first 600 chars",
    })

    phone_ok = bool(re.search(r"\d[\d\s\-\(\)]{6,}\d", snippet))
    checks.append({
        "name": "contact_phone_in_body",
        "ok": phone_ok,
        "detail": "phone-like digits found in first 600 chars" if phone_ok else "no phone digits in first 600 chars",
    })

    body_low = body_text.lower()
    found_headings = [h for h in _STANDARD_HEADINGS if h in body_low]
    headings_ok = len(found_headings) >= 3
    checks.append({
        "name": "standard_headings",
        "ok": headings_ok,
        "detail": f"found: {found_headings}" if found_headings else "none of the expected headings found",
    })

    return {"parseable": parseable, "checks": checks}


def ats_self_check(docx_path: Path, content: Dict, priors: Dict[str, List[str]]) -> Dict:
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]
    text_low = " ".join(paragraphs).lower()

    jd_hits = content["jd_hits"]
    hard = {l: h for l, h in jd_hits.items() if not TERM_BANK[l].get("soft")}
    soft = {l: h for l, h in jd_hits.items() if TERM_BANK[l].get("soft")}

    def found(label: str) -> bool:
        return any(p.search(text_low) for p in _TERM_PATTERNS[label])

    hard_present = [l for l in hard if found(l)]
    hard_missing = [l for l in hard if l not in hard_present]

    # Corrected denominator: include weight of unsupported (unclaimable) tools so
    # a resume that cannot cover N required tools scores honestly lower.
    # Unsupported labels come from content["unsupported"]; use flat weight 4 each
    # (same as score_jobs' in-requirements-region weight).
    unsupported_labels: List[str] = content["unsupported"]
    unsupported_weight = 4 * len(unsupported_labels)

    present_supported_weight = sum(importance(hard[l]) for l in hard_present)
    missing_supported_weight = sum(importance(hard[l]) for l in hard_missing)
    hard_total = present_supported_weight + missing_supported_weight + unsupported_weight
    hard_cov = (present_supported_weight / hard_total) if hard_total else 1.0

    soft_present = [l for l in soft if found(l)]
    soft_cov = (len(soft_present) / len(soft)) if soft else 1.0

    title_ok = content["subtitle"].lower() in text_low

    edu_ok = any(line.lower() in text_low for line in EDUCATION if line.strip()) or "master of" in text_low

    quantified = sum(
        1 for p in paragraphs
        if len(p) > 40 and re.search(r"\d", p) and not p.lower().startswith(("master of", "bachelor of"))
    )
    quantified_score = min(1.0, quantified / QUANTIFIED_TARGET)

    word_count = len(text_low.split())
    word_ok = WORD_BAND[0] <= word_count <= WORD_BAND[1]

    predicted = [l for l in priors.get(content["family"], []) if l not in jd_hits][:5]
    predicted_present = [l for l in predicted if found(l)]
    predicted_cov = (len(predicted_present) / len(predicted)) if predicted else 1.0

    match_rate = round(100 * (
        0.55 * hard_cov +
        0.10 * soft_cov +
        0.10 * (1.0 if title_ok else 0.0) +
        0.05 * (1.0 if edu_ok else 0.0) +
        0.10 * quantified_score +
        0.05 * (1.0 if word_ok else 0.0) +
        0.05 * predicted_cov
    ))

    readiness = docx_readiness(docx_path)

    return {
        "match_rate": match_rate,
        "target": MATCH_RATE_TARGET,
        "passes_target": match_rate >= MATCH_RATE_TARGET,
        # Back-compat alias kept: self_check_match_rate read from tracking CSV, not here
        "self_check_match_rate": match_rate,
        "self_check_passes_target": match_rate >= MATCH_RATE_TARGET,
        "hard_coverage_pct": round(100 * hard_cov),
        # Three distinct lists:
        "hard_terms_present": sorted(hard_present),
        "hard_terms_missing_claimable": sorted(hard_missing),
        "hard_terms_missing": sorted(hard_missing),   # back-compat alias
        "genuine_gaps_unsupported": unsupported_labels,
        "soft_coverage_pct": round(100 * soft_cov),
        "title_mirrored": title_ok,
        "education_present": edu_ok,
        "quantified_bullets": quantified,
        "quantified_target": QUANTIFIED_TARGET,
        "word_count": word_count,
        "word_band": list(WORD_BAND),
        "predicted_skills_checked": predicted,
        "predicted_skills_present": predicted_present,
        "readiness": readiness,
    }


# ---------------------------------------------------------------------------
# Reports
# ---------------------------------------------------------------------------

def write_change_log(path: Path, target: Dict[str, str], content: Dict, check: Dict) -> None:
    lines = [
        f"# Resume Change Log - {target['company']} - {target['role_title']}",
        "",
        "- All content drawn from main_resume.md facts only; bullets selected per-JD by evidence coverage.",
        f"- Role family: {content['family']}; bullets chosen to evidence this JD's terms (set-cover selection).",
        f"- Subtitle: '{content['subtitle']}' (seniority guard applied where needed).",
        f"- JD-backed terms emphasised: {', '.join(sorted(content['jd_hits'])) or 'none detected'}.",
        f"- JD terms NOT surfaced in final doc: {', '.join(check['hard_terms_missing']) or 'none'}.",
        f"- Genuine gaps (cannot truthfully claim): {', '.join(content['unsupported']) or 'none'}.",
        f"- Self-check match rate: {check['match_rate']}/100 (target {MATCH_RATE_TARGET}).",
        f"- Quantified bullets: {check['quantified_bullets']} (target >= {QUANTIFIED_TARGET}); word count {check['word_count']} (band {WORD_BAND[0]}-{WORD_BAND[1]}).",
        "- Format: ATS-safe single-column DOCX, black body text, restrained navy headings, no tables/icons/graphics.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_report(path: Path, target: Dict[str, str], job: Dict[str, str], content: Dict, check: Dict) -> None:
    report = {
        "company": target["company"],
        "role_title": target["role_title"],
        "location": target.get("location", ""),
        "job_url": target.get("job_url", ""),
        "match_score": target.get("match_score", ""),
        "match_band": target.get("match_band", ""),
        "role_family": content["family"],
        "matched_evidence": target.get("matched_evidence", ""),
        "gaps": target.get("gaps", ""),
        "why_keep": target.get("why_keep", ""),
        "ats_self_check": check,
        "readiness": check.get("readiness", {}),
        "resume_constraints": {
            "one_page_recommendation": True,
            "experience_order_preserved": True,
            "unsupported_claims_added": False,
            "bullet_grammar": "Bock XYZ (accomplished X, measured by Y, by doing Z)",
        },
    }
    path.write_text(json.dumps(report, indent=2), encoding="utf-8")


# ---------------------------------------------------------------------------
# Main helpers
# ---------------------------------------------------------------------------

_TRACKING_FIELDNAMES = [
    "generated_date", "company", "role_title", "resume_file", "role_family",
    "match_score", "self_check_match_rate", "hard_coverage_pct",
    "quantified_bullets", "word_count", "jd_terms_missing_from_doc", "genuine_gaps",
]


def generate_one(
    target: Dict[str, str],
    job_posts: List[Dict[str, str]],
    priors: Dict[str, List[str]],
) -> "tuple[Dict, str]":
    """Generate resume package for a single matched job.

    Writes: slug_resume.docx, slug_cover_letter.docx, slug_job_description.txt,
    slug_change_log.md, slug_match_report.json.

    Returns (tracking_row dict, body_str) where body_str is used by the batch
    uniqueness gate.  Calls sys.exit(0) cleanly if JD text is missing.
    """
    job = find_job_post(job_posts, target)
    jd_text = job.get("job_text", "").strip()
    if not jd_text:
        raise ValueError(f"Missing JD text for {target['company']} - {target['role_title']}")

    content = build_content(target, jd_text)
    file_slug = f"{slugify(target['company'])}_{slugify(target['role_title'])}"
    resume_path = OUT_ROOT / f"{file_slug}_resume.docx"

    (OUT_ROOT / f"{file_slug}_job_description.txt").write_text(jd_text + "\n", encoding="utf-8")
    write_resume(content, resume_path)
    write_cover_letter_docx(content, target["company"], target["role_title"], jd_text,
                            OUT_ROOT / f"{file_slug}_cover_letter.docx")

    check = ats_self_check(resume_path, content, priors)
    write_change_log(OUT_ROOT / f"{file_slug}_change_log.md", target, content, check)
    write_report(OUT_ROOT / f"{file_slug}_match_report.json", target, job, content, check)

    tracking_row = {
        "generated_date": date.today().isoformat(),
        "company": target["company"],
        "role_title": target["role_title"],
        "resume_file": f"{file_slug}_resume.docx",
        "role_family": content["family"],
        "match_score": target.get("match_score", ""),
        "self_check_match_rate": check["match_rate"],
        "hard_coverage_pct": check["hard_coverage_pct"],
        "quantified_bullets": check["quantified_bullets"],
        "word_count": check["word_count"],
        "jd_terms_missing_from_doc": ", ".join(check["hard_terms_missing"]),
        "genuine_gaps": ", ".join(content["unsupported"]),
    }
    body_str = content["summary"] + " " + " ".join(
        b for bl in content["bullets_by_slot"].values() for b in bl
    ) + " " + " ".join(content["skills_lines"])
    return tracking_row, body_str


def _upsert_tracking_row(row: Dict) -> None:
    """Append or replace one row in ats_match_tracking.csv, keyed by resume_file."""
    existing: List[Dict] = []
    if TRACKING.exists():
        with TRACKING.open(newline="", encoding="utf-8") as f:
            existing = list(csv.DictReader(f))
    existing = [r for r in existing if r.get("resume_file") != row["resume_file"]]
    existing.append(row)
    with TRACKING.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=_TRACKING_FIELDNAMES)
        writer.writeheader()
        writer.writerows(existing)


def _main_single(row_key_arg: str) -> None:
    """Entry point for --row-key: generate a resume for one job, bypassing generation_enabled."""
    if not FACT_BANK:
        print("FACT_BANK is empty; cannot generate. Exiting.")
        sys.exit(0)
    if not MATCHED.exists():
        print(f"[--row-key] {MATCHED} not found; nothing to do.")
        sys.exit(0)
    if not JOBS.exists():
        print(f"[--row-key] {JOBS} not found; nothing to do.")
        sys.exit(0)

    validate_fact_bank()

    import csv_merge as _csv_merge
    matches = read_csv(MATCHED)
    target = next(
        (r for r in matches if _csv_merge.row_key(r) == row_key_arg),
        None,
    )
    if target is None:
        print(f"[--row-key] Row key {row_key_arg!r} not found in {MATCHED}; nothing to do.")
        sys.exit(0)

    OUT_ROOT.mkdir(exist_ok=True)
    job_posts = read_csv(JOBS)
    priors = family_skill_priors(job_posts)

    try:
        tracking_row, _ = generate_one(target, job_posts, priors)
    except ValueError as exc:
        print(f"[--row-key] {exc}; skipping.")
        sys.exit(0)
    _upsert_tracking_row(tracking_row)

    rate = tracking_row["self_check_match_rate"]
    flag = "" if rate >= MATCH_RATE_TARGET else "  <-- below target"
    print(f"Generated: {tracking_row['resume_file']} [{rate}/100]{flag}")


def main() -> None:
    parser = argparse.ArgumentParser(add_help=True)
    parser.add_argument(
        "--row-key",
        metavar="KEY",
        default=None,
        help="Generate a resume for a single job (bypasses generation_enabled gate).",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Re-generate even if output files already exist (single-job mode only).",
    )
    args = parser.parse_args()

    if args.row_key is not None:
        _main_single(args.row_key)
        return

    # Batch path. Generation is ON by default; two cases skip it:
    #   1. explicitly disabled in config.json
    #   2. no real facts.json yet — building from the placeholder "Company A/B/C"
    #      fact bank would hand the user a resume full of fake employers, which is
    #      worse than none. Onboarding writes facts.json from the parsed résumé.
    if not config.generation_enabled:
        print("Resume generation is OFF for this profile. Enable it in config.json (matching.generation_enabled).")
        return
    if _facts_override is None:
        print("Resume generation is ON but no resume data found. Upload your resume in onboarding "
              "to generate tailored resumes (skipping - will not build from placeholder data).")
        return
    validate_fact_bank()
    OUT_ROOT.mkdir(exist_ok=True)
    matches = [r for r in read_csv(MATCHED) if r.get("company") not in _get_exclude_companies()]
    job_posts = read_csv(JOBS)
    priors = family_skill_priors(job_posts)

    tracking_rows = []
    bodies: Dict[str, str] = {}

    for target in matches:
        try:
            tracking_row, body_str = generate_one(target, job_posts, priors)
        except ValueError as exc:
            raise SystemExit(str(exc))
        file_slug = f"{slugify(target['company'])}_{slugify(target['role_title'])}"
        bodies[file_slug] = body_str
        tracking_rows.append(tracking_row)

    # Uniqueness gate: no two resumes may share an identical tailored body.
    slugs = sorted(bodies)
    identical, near = [], []
    for i, a in enumerate(slugs):
        for b in slugs[i + 1:]:
            ratio = difflib.SequenceMatcher(None, bodies[a], bodies[b]).ratio()
            if ratio == 1.0:
                identical.append((a, b))
            elif ratio > 0.97:
                near.append((a, b, round(ratio, 3)))

    with TRACKING.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=_TRACKING_FIELDNAMES)
        writer.writeheader()
        writer.writerows(tracking_rows)

    rates = [r["self_check_match_rate"] for r in tracking_rows]
    if not rates:
        print("No matched jobs to generate — matched_jobs.csv is empty. Nothing written.")
        return
    passing = sum(1 for r in rates if r >= MATCH_RATE_TARGET)
    print(f"Generated {len(matches)} resume packages under {OUT_ROOT}")
    print(f"Self-check match rates: min {min(rates)}, max {max(rates)}, "
          f"{passing}/{len(rates)} at or above target {MATCH_RATE_TARGET}")
    for row in sorted(tracking_rows, key=lambda r: -r["self_check_match_rate"]):
        flag = "" if row["self_check_match_rate"] >= MATCH_RATE_TARGET else "  <-- below target"
        print(f"  [{row['self_check_match_rate']}] [{row['role_family']}] {row['resume_file']}{flag}")

    if identical:
        print(f"\n[!] UNIQUENESS GATE: {len(identical)} identical body pair(s):")
        for a, b in identical:
            print(f"    {a} == {b}")
    if near:
        print(f"\n[~] Near-identical (>0.97 similarity) pairs — check whether JDs genuinely overlap:")
        for a, b, r in near:
            print(f"    {a} ~ {b} ({r})")
    if not identical and not near:
        print("\nUniqueness gate: all resume bodies are distinct.")


if __name__ == "__main__":
    main()

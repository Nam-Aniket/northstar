"""Standalone résumé builder — python-docx only, no pipeline imports."""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


_NAVY = RGBColor(31, 78, 121)


def _add_bottom_border(paragraph) -> None:
    """Add a thin navy bottom border to a paragraph (ATS-safe: pPr/pBdr technique)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")        # 0.75 pt
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _NAVY
    _add_bottom_border(p)


def _set_font(run, size_pt: float, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold


def build_resume_docx(data: dict) -> Document:
    """
    Build and return a python-docx Document from the résumé data dict.

    data shape:
        name, email, phone, location, linkedin, summary,
        skills (comma-separated str),
        experiences: [{role, company, dates, bullets: [str, ...]}],
        education:   [{degree, school, year}]
    """
    doc = Document()

    # ── Page setup: A4, ~2 cm margins ────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2.0))

    # Default paragraph style: Calibri 10.5 pt, line spacing 1.05
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = Pt(10.5 * 1.05)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)

    # ── Name ─────────────────────────────────────────────────────────────────
    name = (data.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        _set_font(run, 18, bold=True)
        run.font.color.rgb = _NAVY

    # ── Contact line ─────────────────────────────────────────────────────────
    contact_parts = [
        data.get("email", "").strip(),
        data.get("phone", "").strip(),
        data.get("location", "").strip(),
        data.get("linkedin", "").strip(),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    if contact_line:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(contact_line)
        _set_font(run, 9)
        run.font.color.rgb = RGBColor(87, 83, 74)

    # ── Summary ──────────────────────────────────────────────────────────────
    summary = (data.get("summary") or "").strip()
    if summary:
        _section_heading(doc, "Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)

    # ── Experience ───────────────────────────────────────────────────────────
    experiences = data.get("experiences") or []
    experiences = [e for e in experiences if isinstance(e, dict)]
    if experiences:
        _section_heading(doc, "Experience")
        for exp in experiences:
            role = (exp.get("role") or "").strip()
            company = (exp.get("company") or "").strip()
            dates = (exp.get("dates") or "").strip()
            bullets = [b.strip() for b in (exp.get("bullets") or []) if b.strip()]

            if not (role or company or bullets):
                continue

            # Role — Company · Dates on one paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)

            if role:
                r = p.add_run(role)
                _set_font(r, 10.5, bold=True)

            if company:
                sep = " — " if role else ""
                r = p.add_run(sep + company)
                _set_font(r, 10.5)

            if dates:
                r = p.add_run("  " + dates)
                _set_font(r, 9.5)
                r.font.color.rgb = RGBColor(87, 83, 74)

            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Cm(0.5)
                r = bp.add_run(bullet)
                _set_font(r, 10.5)

    # ── Skills ───────────────────────────────────────────────────────────────
    skills = (data.get("skills") or "").strip()
    if skills:
        _section_heading(doc, "Skills")
        p = doc.add_paragraph(skills)
        p.paragraph_format.space_after = Pt(4)

    # ── Education ────────────────────────────────────────────────────────────
    education = data.get("education") or []
    education = [e for e in education if isinstance(e, dict)]
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            degree = (edu.get("degree") or "").strip()
            school = (edu.get("school") or "").strip()
            year = (edu.get("year") or "").strip()
            parts = [p for p in [degree, school, year] if p]
            if parts:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(" — ".join(parts))
                _set_font(r, 10.5)

    return doc
